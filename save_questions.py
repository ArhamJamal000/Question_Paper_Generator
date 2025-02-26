import docx
from docx.shared import Pt, Inches

def save_question_paper(saqs, laqs, file_path, course_name, course_code, examination_date, exam_type):
    print("SAQs:")
    for saq in saqs:
        print(saq['question'])
    print("LAQs:")
    for laq in laqs:
        print(laq['question'])

    document = docx.Document()

    # Set default font
    style = document.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(12)

    # Add course code
    course_code_paragraph = document.add_paragraph()
    course_code_paragraph.alignment = 2  # Right alignment
    course_code_run = course_code_paragraph.add_run(f"Course code: {course_code}")
    course_code_run.font.size = Pt(12)
    course_code_run.font.bold = True

    # Add the logo
    logo_paragraph = document.add_paragraph()
    logo_paragraph.alignment = 1  # Center alignment
    logo_run = logo_paragraph.add_run()
    logo_run.add_picture('logo.png', width=Inches(0.582), height=Inches(0.708))  # Update with the correct path

    # Institute name and details
    institute_name_paragraph = document.add_paragraph()
    institute_name_paragraph.alignment = 1  # Center alignment
    institute_name_run = institute_name_paragraph.add_run(
        "LORDS INSTITUTE OF ENGINEERING AND TECHNOLOGY \n (UGC AUTONOMOUS)"
    )
    institute_name_run.font.size = Pt(14)
    institute_name_run.font.bold = True

    institute_details = [
        "Approved by AICTE | Recognized by Government of Telangana | Affiliated to Osmania University",
        "Accredited by NBA | Accredited with ‘A’ grade by NAAC | Accredited by NABL"
    ]
    for detail in institute_details:
        detail_paragraph = document.add_paragraph(detail)
        detail_paragraph.alignment = 1  # Center alignment
        detail_paragraph.runs[0].font.size = Pt(10)
        detail_paragraph.runs[0].font.bold = True

    # Course details
    course_details_paragraph = document.add_paragraph()
    if exam_type == "SEE":
        course_details_paragraph_run = course_details_paragraph.add_run(
            f"B.E - SEMESTER END EXAMINATION QUESTION PAPER {examination_date} "
        )
    elif exam_type == "CIE":
        course_details_paragraph_run = course_details_paragraph.add_run(
            f"B.E - CONTINUOUS INTERNAL EVALUATION QUESTION PAPER {examination_date} "
        )
    else:
        course_details_paragraph_run = course_details_paragraph.add_run(
            f"B.E ({exam_type}) QUESTION PAPER {examination_date} ({exam_type})"
        )
    course_details_paragraph_run.font.size = Pt(12)
    course_details_paragraph_run.font.bold = True
    course_details_paragraph.alignment = 1  # Center alignment

    # Course name
    course_name_paragraph = document.add_paragraph()
    course_name_run = course_name_paragraph.add_run(f"Course name: {course_name}")
    course_name_run.font.size = Pt(12)
    course_name_run.font.bold = True
    course_name_paragraph.alignment = 2  # Right alignment

    # Time and Max Marks
    if exam_type == "SEE":
        time_paragraph = document.add_paragraph("Time: 3 Hours")
        marks_paragraph = document.add_paragraph("Max. Marks : 60")
    elif exam_type == "CIE":
        time_paragraph = document.add_paragraph("Time: 1 Hour")
        marks_paragraph = document.add_paragraph("Max. Marks : 20")
    else:
        time_paragraph = document.add_paragraph("Time: __ Hours")
        marks_paragraph = document.add_paragraph("Max. Marks : __")
    time_paragraph.alignment = 0  # Left alignment
    time_paragraph.runs[0].font.size = Pt(12)
    time_paragraph.runs[0].font.bold = True
    marks_paragraph.alignment = 2  # Right alignment
    marks_paragraph.runs[0].font.size = Pt(12)
    marks_paragraph.runs[0].font.bold = True

    # Bloom's Taxonomy Levels
    taxonomy_paragraph = document.add_paragraph("Bloom’s Taxonomy Levels (BTL)")
    taxonomy_paragraph.runs[0].font.bold = True
    taxonomy_list = ["I. Remember", "II. Understand", "III.  Apply", "IV.  Analyze", "V.  Evaluate", "VI. Create"]

    taxonomy_table = document.add_table(rows=1, cols=6)
    for idx, level in enumerate(taxonomy_list):
        cell = taxonomy_table.cell(0, idx)
        cell.text = level
        cell.paragraphs[0].alignment = 1  # Center alignment
        cell.paragraphs[0].runs[0].font.bold = True

    # Instructions
    instructions_paragraph = document.add_paragraph("Instructions to the Students:")
    instructions_paragraph.alignment = 0  # Left alignment
    instructions_run = instructions_paragraph.runs[0]
    instructions_run.font.size = Pt(12)
    instructions_run.font.bold = True

    instruction_list = []
    if exam_type == "SEE":
        instruction_list = [
            "• Question No. 1 is compulsory",
            "• Answer any 4 questions from Q.No.2 – Q.No7"
        ]
    elif exam_type == "CIE":
        instruction_list = [
            "• Question No. 1 is compulsory",
            "• Answer any 2 questions from Q.No.2 – Q.No4"
        ]
    else:
        instruction_list = [
            "• Question No. 1 is compulsory"
        ]

    for instruction in instruction_list:
        instruction_paragraph = document.add_paragraph(instruction)
        instruction_paragraph.alignment = 0  # Left alignment
        instruction_paragraph.paragraph_format.left_indent = Inches(0.25)
        instruction_paragraph.runs[0].font.size = Pt(12)

    # SAQs Section
    if saqs:
        saq_paragraph = document.add_paragraph()
        saq_paragraph.alignment = 0  # Left alignment
        if exam_type == "SEE":
            saq_run = saq_paragraph.add_run(
                "Q.No. 1: Five short answer Questions covering entire content of 5 units of the syllabus. At least one question from each Unit."
            )
        elif exam_type == "CIE":
            saq_run = saq_paragraph.add_run(
                "Q.No. 1: Three short answer Questions covering entire content of 3 units of the syllabus. At least one question from each Unit."
            )
        else:{}
           #saq_run = saq_paragraph.add_run(
                #"Q.No. 1: Short answer Questions covering the entire content of the syllabus.")
        saq_run.font.size = Pt(12)
        saq_run.font.bold = True

        # Create SAQ table
        saq_table = document.add_table(rows=1, cols=5)
        saq_table.style = 'Table Grid'
        saq_table.autofit = False
        saq_table.columns[0].width = Inches(0.5)
        saq_table.columns[1].width = Inches(2)  # Set the second column to 4 times bigger
        saq_table.columns[2].width = Inches(0.5)
        saq_table.columns[3].width = Inches(0.5)
        saq_table.columns[4].width = Inches(0.5)

        hdr_cells = saq_table.rows[0].cells
        hdr_cells[0].text = ''
        hdr_cells[1].text = ''
        hdr_cells[2].text = 'Marks'
        hdr_cells[3].text = 'CO'
        hdr_cells[4].text = 'BTL'

        for i, question in enumerate(saqs, start=1):
            row_cells = saq_table.add_row().cells
            row_cells[0].text = chr(96 + i)  # a, b, c, etc.
            row_cells[1].text = f"{chr(97 + (i % 2))}.{question['question']}"
            row_cells[2].text = str(question['marks'])
            row_cells[3].text = f"CO{question['btl']}"
            row_cells[4].text = f"BTL{question['co']}"



    # Create LAQ table
    laq_table = document.add_table(rows=1, cols=5)
    laq_table.style = 'Table Grid'
    laq_table.autofit = False
    laq_table.columns[0].width = Inches(0.5)
    laq_table.columns[1].width = Inches(2)  # Set the second column to 4 times bigger
    laq_table.columns[2].width = Inches(0.5)
    laq_table.columns[3].width = Inches(0.5)
    laq_table.columns[4].width = Inches(0.5)

    hdr_cells = laq_table.rows[0].cells
    hdr_cells[0].text = ''
    hdr_cells[1].text = ''
    hdr_cells[2].text = 'Marks'
    hdr_cells[3].text = 'CO'
    hdr_cells[4].text = 'BTL'

    if exam_type == "CIE":
        for i in range(0, len(laqs)):
            question_a = laqs[i]
            row_cells = laq_table.add_row().cells
            row_cells[0].text = f"{i // 2 + 2}."  # Question numbering
            row_cells[1].text = question_a['question']
            row_cells[2].text = str(question_a['marks'])
            row_cells[3].text = f"CO{question_a['btl']}"
            row_cells[4].text = f"BTL{question_a['co']}"





    else:
        for i in range(0, len(laqs), 2):
            question_a = laqs[i]
            row_cells = laq_table.add_row().cells
            row_cells[0].text = f"{i // 2 + 2}.a"  # Question numbering
            row_cells[1].text = question_a['question']
            row_cells[2].text = str(question_a['marks'])
            row_cells[3].text = f"CO{question_a['btl']}"
            row_cells[4].text = f"BTL{question_a['co']}"

            if i + 1 < len(laqs):
                question_b = laqs[i + 1]
                row_cells_b = laq_table.add_row().cells
                row_cells_b[0].text = "   b"
                row_cells_b[1].text = question_b['question']
                row_cells_b[2].text = str(question_b['marks'])
                row_cells_b[3].text = f"CO{question_b['btl']}"
                row_cells_b[4].text = f"BTL{question_b['btl']}"


    # Footer
    footer_paragraph = document.add_paragraph("\n\n**************")
    footer_paragraph.alignment = 1  # Center alignment
    footer_run = footer_paragraph.runs[0]
    footer_run.font.size = Pt(12)
    footer_run.font.bold = True

    # Save document
    document.save(file_path)