from docx import Document

# Create a new Word document
doc = Document()

# Define data for each unit
questions_by_unit = {
    "Unit-1": [
        ["U-1 SAQ-1", 2, "BTL-1",1], ["U-1 SAQ-2", 2, "BTL-2",5], ["U-1 SAQ-3", 2, "BTL-3",1], ["U-1 SAQ-4", 2, "BTL-1",2],
        ["U-1 SAQ-5", 2, "BTL-2",2], ["U-1 SAQ-6", 2, "BTL-3",4], ["U-1 SAQ-7", 2, "BTL-1",2], ["U-1 SAQ-8", 2, "BTL-2",1],
        ["U-1 SAQ-9", 2, "BTL-3",3], ["U-1 SAQ-10", 2, "BTL-1",3], ["U-1 LAQ-1", 6, "BTL-4",3], ["U-1 LAQ-2", 6, "BTL-5",3],
        ["U-1 LAQ-3", 6, "BTL-6",4], ["U-1 LAQ-4", 6, "BTL-4",2], ["U-1 LAQ-5", 6, "BTL-5",4], ["U-1 LAQ-6", 6, "BTL-6",4],
        ["U-1 LAQ-7", 6, "BTL-4",5], ["U-1 LAQ-8", 6, "BTL-5",1], ["U-1 LAQ-9", 6, "BTL-6",5], ["U-1 LAQ-10", 6, "BTL-4",5]
    ],
    "Unit-2": [
        ["U-2 SAQ-1", 2, "BTL-2"], ["U-2 SAQ-2", 2, "BTL-3"], ["U-2 SAQ-3", 2, "BTL-1"], ["U-2 SAQ-4", 2, "BTL-4"],
        ["U-2 SAQ-5", 2, "BTL-5"], ["U-2 SAQ-6", 2, "BTL-2"], ["U-2 SAQ-7", 2, "BTL-3"], ["U-2 SAQ-8", 2, "BTL-1"],
        ["U-2 SAQ-9", 2, "BTL-4"], ["U-2 SAQ-10", 2, "BTL-5"], ["U-2 LAQ-1", 6, "BTL-4"], ["U-2 LAQ-2", 6, "BTL-5"],
        ["U-2 LAQ-3", 6, "BTL-6"], ["U-2 LAQ-4", 6, "BTL-4"], ["U-2 LAQ-5", 6, "BTL-5"], ["U-2 LAQ-6", 6, "BTL-6"],
        ["U-2 LAQ-7", 6, "BTL-4"], ["U-2 LAQ-8", 6, "BTL-5"], ["U-2 LAQ-9", 6, "BTL-6"], ["U-2 LAQ-10", 6, "BTL-4"]
    ],
    "Unit-3": [
        ["U-3 SAQ-1", 2, "BTL-1"], ["U-3 SAQ-2", 2, "BTL-2"], ["U-3 SAQ-3", 2, "BTL-3"], ["U-3 SAQ-4", 2, "BTL-1"],
        ["U-3 SAQ-5", 2, "BTL-2"], ["U-3 SAQ-6", 2, "BTL-3"], ["U-3 SAQ-7", 2, "BTL-1"], ["U-3 SAQ-8", 2, "BTL-2"],
        ["U-3 SAQ-9", 2, "BTL-3"], ["U-3 SAQ-10", 2, "BTL-1"], ["U-3 LAQ-1", 6, "BTL-4"], ["U-3 LAQ-2", 6, "BTL-5"],
        ["U-3 LAQ-3", 6, "BTL-6"], ["U-3 LAQ-4", 6, "BTL-4"], ["U-3 LAQ-5", 6, "BTL-5"], ["U-3 LAQ-6", 6, "BTL-6"],
        ["U-3 LAQ-7", 6, "BTL-4"], ["U-3 LAQ-8", 6, "BTL-5"], ["U-3 LAQ-9", 6, "BTL-6"], ["U-3 LAQ-10", 6, "BTL-4"]
    ],
    "Unit-4": [
        ["U-4 SAQ-1", 2, "BTL-1"], ["U-4 SAQ-2", 2, "BTL-2"], ["U-4 SAQ-3", 2, "BTL-3"], ["U-4 SAQ-4", 2, "BTL-1"],
        ["U-4 SAQ-5", 2, "BTL-2"], ["U-4 SAQ-6", 2, "BTL-3"], ["U-4 SAQ-7", 2, "BTL-1"], ["U-4 SAQ-8", 2, "BTL-2"],
        ["U-4 SAQ-9", 2, "BTL-3"], ["U-4 SAQ-10", 2, "BTL-1"], ["U-4 LAQ-1", 6, "BTL-4"], ["U-4 LAQ-2", 6, "BTL-5"],
        ["U-4 LAQ-3", 6, "BTL-6"], ["U-4 LAQ-4", 6, "BTL-4"], ["U-4 LAQ-5", 6, "BTL-5"], ["U-4 LAQ-6", 6, "BTL-6"],
        ["U-4 LAQ-7", 6, "BTL-4"], ["U-4 LAQ-8", 6, "BTL-5"], ["U-4 LAQ-9", 6, "BTL-6"], ["U-4 LAQ-10", 6, "BTL-4"]
    ],
    "Unit-5": [
        ["U-5 SAQ-1", 2, "BTL-1"], ["U-5 SAQ-2", 2, "BTL-2"], ["U-5 SAQ-3", 2, "BTL-3"], ["U-5 SAQ-4", 2, "BTL-1"],
        ["U-5 SAQ-5", 2, "BTL-2"], ["U-5 SAQ-6", 2, "BTL-3"], ["U-5 SAQ-7", 2, "BTL-1"], ["U-5 SAQ-8", 2, "BTL-2"],
        ["U-5 SAQ-9", 2, "BTL-3"], ["U-5 SAQ-10", 2, "BTL-1"], ["U-5 LAQ-1", 6, "BTL-4"], ["U-5 LAQ-2", 6, "BTL-5"],
        ["U-5 LAQ-3", 6, "BTL-6"], ["U-5 LAQ-4", 6, "BTL-4"], ["U-5 LAQ-5", 6, "BTL-5"], ["U-5 LAQ-6", 6, "BTL-6"],
        ["U-5 LAQ-7", 6, "BTL-4"], ["U-5 LAQ-8", 6, "BTL-5"], ["U-5 LAQ-9", 6, "BTL-6"], ["U-5 LAQ-10", 6, "BTL-4"]
    ]
}

# Create 5 tables for each unit
# Create 5 tables for each unit
for unit, questions in questions_by_unit.items():
    doc.add_paragraph(unit)  # Add unit heading
    table = doc.add_table(rows=1, cols=4)
    table.style = 'Table Grid'

    # Add headers
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Questions'
    hdr_cells[1].text = 'Marks'
    hdr_cells[2].text = 'CO'
    hdr_cells[3].text = 'BTL Level'

    # Add data to the table (20 SAQs and 20 LAQs)
    for question in questions:
        row_cells = table.add_row().cells
        row_cells[0].text = question[0]
        row_cells[1].text = str(question[1])
        row_cells[2].text = 'CO' + question[2][3:]  # Extract CO from BTL Level
        row_cells[3].text = question[2]

# Save the document

doc.save("question_paper1.docx")