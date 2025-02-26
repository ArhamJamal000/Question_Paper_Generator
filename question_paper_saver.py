from docx import Document
import docx
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement


from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt

def add_question_analysis_table(document, questions):
    """
    Add Bloom's Taxonomy Level (BTL) analysis table with counts and percentages to a new page in the document.
    Includes additional rows for SAQ BTL1 percentage, SAQ BTL2 percentage, LAQ BTL2 percentage, and BTL3-BTL6 percentage.
    """
    # Add a new page
    document.add_page_break()

    # Add heading
    analysis_heading = document.add_paragraph("Question Analysis")
    analysis_heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
    analysis_heading.runs[0].font.bold = True
    analysis_heading.runs[0].font.size = Pt(14)

    # Create the analysis table
    table = document.add_table(rows=7, cols=6)
    table.style = 'Table Grid'

    # Populate the header row
    headers = ["BTL-1 (Remember)", "BTL-2 (Understand)", "BTL-3 (Apply)",
               "BTL-4 (Analyze)", "BTL-5 (Evaluate)", "BTL-6 (Create)"]
    header_cells = table.rows[0].cells
    for i, header in enumerate(headers):
        cell = header_cells[i]
        cell.text = header
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        cell.paragraphs[0].runs[0].font.bold = True

    # Analyze the questions and count BTL levels
    btl_counts = [0] * 6  # Initialize counts for BTL-1 to BTL-6
    saq_counts = [0] * 6  # SAQ counts for BTL-1 to BTL-6
    laq_counts = [0] * 6  # LAQ counts for BTL-1 to BTL-6

    for question in questions:
        btl_level = int(question.get('btl'))  # BTL level should be in the question dictionary
        question_type = question.get('type', '').lower()  # Question type (e.g., 'saq', 'laq')
        if 1 <= btl_level <= 6:
            btl_counts[btl_level - 1] += 1
            if question_type == 'saq':
                saq_counts[btl_level - 1] += 1
            elif question_type == 'laq':
                laq_counts[btl_level - 1] += 1

    # Add counts to the second row
    count_cells = table.rows[1].cells
    total_questions = len(questions)
    for i, count in enumerate(btl_counts):
        count_cells[i].text = str(count)
        count_cells[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Calculate percentages and add to the third row
    percentage_cells = table.rows[2].cells
    for i, count in enumerate(btl_counts):
        percentage = (count / total_questions) * 100 if total_questions > 0 else 0
        percentage_cells[i].text = f"{percentage:.2f}%"  # Format percentage to 2 decimal places
        percentage_cells[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Add SAQ BTL1 percentage (row 4)
    saq_btl1_percentage = (saq_counts[0] / total_questions) * 100 if total_questions > 0 else 0
    row_4_cells = table.rows[3].cells
    row_4_cells[0].merge(row_4_cells[-1])  # Merge cells for the label
    row_4_cells[0].text = f"SAQ BTL1 Percentage: {saq_btl1_percentage:.2f}%"
    row_4_cells[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Add SAQ BTL2 percentage (row 5)
    saq_btl2_percentage = (saq_counts[1] / total_questions) * 100 if total_questions > 0 else 0
    row_5_cells = table.rows[4].cells
    row_5_cells[0].merge(row_5_cells[-1])
    row_5_cells[0].text = f"SAQ BTL2 Percentage: {saq_btl2_percentage:.2f}%"
    row_5_cells[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Add LAQ BTL2 percentage (row 6)
    laq_btl2_percentage = (laq_counts[1] / total_questions) * 100 if total_questions > 0 else 0
    row_6_cells = table.rows[5].cells
    row_6_cells[0].merge(row_6_cells[-1])
    row_6_cells[0].text = f"LAQ BTL2 Percentage: {laq_btl2_percentage:.2f}%"
    row_6_cells[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Add BTL3 to BTL6 percentage (row 7)
    btl3_to_btl6_count = sum(btl_counts[2:])
    btl3_to_btl6_percentage = (btl3_to_btl6_count / total_questions) * 100 if total_questions > 0 else 0
    row_7_cells = table.rows[6].cells
    row_7_cells[0].merge(row_7_cells[-1])
    row_7_cells[0].text = f"BTL3 to BTL6 Percentage: {btl3_to_btl6_percentage:.2f}%"
    row_7_cells[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
def remove_table_borders(table):
    """Remove all borders from the given table."""
    tbl = table._element
    tbl_pr = tbl.tblPr
    borders = tbl_pr.xpath(".//w:tblBorders")
    for border in borders:
        tbl_pr.remove(border)



def get_column_widths(table):
    """Calculates the width of each column in a table."""
    column_widths = []
    for row in table.rows:
        for cell in row.cells:
            width = cell.width
            column_index = row.cells.index(cell)
            if len(column_widths) <= column_index:
                column_widths.append(width)
            else:
                column_widths[column_index] = max(column_widths[column_index], width)
    return column_widths


def autofit_column_width(table):
    """Adjust column width based on the content length."""
    for col in table.columns:
        max_length = 0
        # Find the maximum length in each column
        for cell in col.cells:
            max_length = max(max_length, len(cell.text))

        # Set the column width based on max length
        col.width = Inches(max_length * 0.08)  # Adjust multiplier as needed



def save_question_paper(saqs, laqs, file_path, course_name, course_code, examination_date, exam_type):
    print("SAQs:")
    for saq in saqs:
        print(saq['question'])
    print("LAQs:")
    for laq in laqs:
        print(laq['question'])

    document = docx.Document()

    # Set default font
    style = document.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(12)

    # Add course code
    course_code_paragraph = document.add_paragraph()
    course_code_paragraph.alignment = 2  # Right alignment
    course_code_run = course_code_paragraph.add_run(f"Course code: {course_code}")
    course_code_run.font.size = Pt(12)
    course_code_run.font.bold = True

    # Add the logo
    logo_paragraph = document.add_paragraph()
    logo_paragraph.alignment = 1  # Center alignment
    logo_run = logo_paragraph.add_run()
    logo_run.add_picture('logo.png', width=Inches(0.582), height=Inches(0.708))  # Update with the correct path

    # Institute name and details
    institute_name_paragraph = document.add_paragraph()
    institute_name_paragraph.alignment = 1  # Center alignment
    institute_name_run = institute_name_paragraph.add_run(
        "LORDS INSTITUTE OF ENGINEERING AND TECHNOLOGY \n (UGC AUTONOMOUS)"
    )
    institute_name_run.font.size = Pt(14)
    institute_name_run.font.bold = True
    institute_name_run.font.name = 'Times New Roman'  # Set font name

    institute_details = [
        "Approved by AICTE | Recognized by Government of Telangana | Affiliated to Osmania University",
        "Accredited by NBA | Accredited with ‘A’ grade by NAAC | Accredited by NABL"
    ]
    for detail in institute_details:
        detail_paragraph = document.add_paragraph(detail)
        detail_paragraph.alignment = 1  # Center alignment
        run = detail_paragraph.runs[0]  # Access the first run
        run.font.name = 'Times New Roman'  # Set font name
        run.font.size = Pt(10)  # Set font size
        run.font.bold = True  # Set bold

    # Course details
    course_details_paragraph = document.add_paragraph()
    if exam_type == "SEE":
        course_details_paragraph_run = course_details_paragraph.add_run(
            f"B.E - SEMESTER END EXAMINATION QUESTION PAPER {examination_date} "
        )
    elif exam_type == "CIE":
        course_details_paragraph_run = course_details_paragraph.add_run(
            f"B.E - CONTINUOUS INTERNAL EVALUATION QUESTION PAPER {examination_date} "
        )
    else:
        course_details_paragraph_run = course_details_paragraph.add_run(
            f"B.E ({exam_type}) QUESTION PAPER {examination_date} ({exam_type})"
        )
    course_details_paragraph_run.font.size = Pt(12)
    course_details_paragraph_run.font.bold = True
    course_details_paragraph.alignment = 1  # Center alignment
    course_details_paragraph_run.font.name = 'Book Antiqua'

    # Course name
    course_name_paragraph = document.add_paragraph()
    course_name_run = course_name_paragraph.add_run(f"Course name: {course_name}")
    course_name_run.font.size = Pt(14)
    course_name_run.font.bold = True
    course_name_run.font.name = 'Book Antiqua'
    course_name_paragraph.alignment = 1  # Right alignment

    # Time and Max Marks
    if exam_type == "SEE":
        time_paragraph = document.add_paragraph("Time: 3 Hours")
        marks_paragraph = document.add_paragraph("Max. Marks : 60")
    elif exam_type == "CIE":
        time_paragraph = document.add_paragraph("Time: 1 Hour")
        marks_paragraph = document.add_paragraph("Max. Marks : 20")
    else:
        time_paragraph = document.add_paragraph("Time: __ Hours")
        marks_paragraph = document.add_paragraph("Max. Marks : __")
    time_paragraph.alignment = 0  # Left alignment
    time_paragraph.runs[0].font.size = Pt(12)
    time_paragraph.runs[0].font.bold = True
    time_paragraph.runs[0].font.name = 'Book Antiqua'

    marks_paragraph.alignment = 2  # Right alignment
    marks_paragraph.runs[0].font.size = Pt(12)
    marks_paragraph.runs[0].font.bold = True
    marks_paragraph.runs[0].font.name = 'Book Antiqua'

    # Bloom's Taxonomy Levels


    # Taxonomy paragraph - Set to Times New Roman font, bold, and font size 12
    taxonomy_paragraph = document.add_paragraph("Bloom’s Taxonomy Levels (BTL)")
    taxonomy_paragraph.runs[0].font.bold = False
    taxonomy_paragraph.runs[0].font.name = 'Times New Roman'  # Set font to Times New Roman
    taxonomy_paragraph.runs[0].font.size = Pt(12)  # Set font size to 12 pt

    # Taxonomy list
    taxonomy_list = ["I. Remember", "II. Understand", "III.  Apply", "IV.  Analyze", "V.  Evaluate", "VI. Create"]

    # Create the table
    taxonomy_table = document.add_table(rows=1, cols=6)
    for idx, level in enumerate(taxonomy_list):
        cell = taxonomy_table.cell(0, idx)
        cell.text = level
        cell.paragraphs[0].alignment = 1  # Center alignment
        cell.paragraphs[0].runs[0].font.bold = True
        cell.paragraphs[0].runs[0].font.name = 'Times New Roman'  # Set font to Times New Roman
        cell.paragraphs[0].runs[0].font.size = Pt(11)  # Set font size to 11 pt for table elements



    # Instructions paragraph - Set to Book Antiqua font, bold, and font size 12
    instructions_paragraph = document.add_paragraph("Instructions to the Students:")
    instructions_paragraph.alignment = 0  # Left alignment
    instructions_run = instructions_paragraph.runs[0]
    instructions_run.font.size = Pt(12)
    instructions_run.font.bold = True
    instructions_run.font.name = 'Book Antiqua'  # Set font to Book Antiqua

    # Instruction list
    instruction_list = []
    if exam_type == "SEE":
        instruction_list = [
            "• Question No. 1 is compulsory",
            "• Answer any 4 questions from Q.No.2 – Q.No7"
        ]
    elif exam_type == "CIE":
        instruction_list = [
            "• Question No. 1 is compulsory",
            "• Answer any 2 questions from Q.No.2 – Q.No4"
        ]
    else:
        instruction_list = [
            "• Question No. 1 is compulsory"
        ]

    # Add each instruction with Book Antiqua font and size 12
    for instruction in instruction_list:
        instruction_paragraph = document.add_paragraph(instruction)
        instruction_paragraph.alignment = 0  # Left alignment
        instruction_paragraph.paragraph_format.left_indent = Inches(0.25)
        instruction_paragraph.runs[0].font.size = Pt(12)
        instruction_paragraph.runs[0].font.name = 'Book Antiqua'  # Set font to Book Antiqua

    combined_table = document.add_table(rows=1, cols=5)
    combined_table.style = None
    combined_table.autofit = False
    combined_table.columns[0].width = Inches(0.5)
    combined_table.columns[1].width = Inches(2)  # Set the second column to 4 times bigger
    combined_table.columns[2].width = Inches(0.5)
    combined_table.columns[3].width = Inches(0.5)
    combined_table.columns[4].width = Inches(0.5)

    hdr_cells = combined_table.rows[0].cells
    hdr_cells[0].text = ''
    hdr_cells[1].text = ''
    hdr_cells[2].text = ''
    hdr_cells[3].text = 'CO'
    hdr_cells[4].text = 'BTL'

    # SAQs Section
    if saqs:



        # Create SAQ table


        for i, question in enumerate(saqs, start=1):
            row_cells = combined_table.add_row().cells
            if i ==1:
                row_cells[0].text = f"{i}"  # a, b, c, etc.
            else:
                row_cells[0].text = " "
            row_cells[1].text = f"{chr(96 + i)}.{question['question']}"
            row_cells[2].text = f"[{str(question['marks'])}]"
            row_cells[3].text = f"CO{question['co']}"
            row_cells[4].text = f"BTL{question['btl']}"

    qno = 2
    if exam_type == "CIE":
        for i in range(0, len(laqs)):
            question_a = laqs[i]
            row_cells = combined_table.add_row().cells
            row_cells[0].text = f"{qno}"  # Question numbering
            row_cells[1].text = question_a['question']
            row_cells[2].text = f"[7]"
            row_cells[3].text = f"CO{question_a['co']}"
            row_cells[4].text = f"BTL{question_a['btl']}"
            qno+=1





    else:


        qno = 2  # Initialize question number

        for laq in laqs:
            if laq['marks'] == 12:
                # For 12 marks questions, add it once
                row_cells = combined_table.add_row().cells
                row_cells[0].text = f"{qno}"  # Question numbering
                row_cells[1].text = f"{laq['question']}"  # Single entry for 12 marks
                row_cells[2].text = f"[{laq['marks']}]"
                row_cells[3].text = f"CO{laq['co']}"
                row_cells[4].text = f"BTL{laq['btl']}"
                qno += 1  # Increment question number

            elif laq['marks'] == 6:
                # For 6 marks questions, add it as a pair (a and b)
                row_cells = combined_table.add_row().cells
                row_cells[0].text = f"{qno}"  # Question numbering
                row_cells[1].text = f"a. {laq['question']}"  # First part of the pair
                row_cells[2].text = f"[{laq['marks']}]"
                row_cells[3].text = f"CO{laq['co']}"
                row_cells[4].text = f"BTL{laq['btl']}"

                # Check if there's a next question to pair with
                if qno < len(laqs) - 1 and laqs[qno]['marks'] == 6:
                    question_b = laqs[qno]  # Get the next question
                    row_cells_b = combined_table.add_row().cells
                    row_cells_b[0].text = "  "  # Empty for the second part
                    row_cells_b[1].text = f"b. {question_b['question']}"  # Second part of the pair
                    row_cells_b[2].text = f"[{question_b['marks']}]"
                    row_cells_b[3].text = f"CO{question_b['co']}"
                    row_cells_b[4].text = f"BTL{question_b['btl']}"
                    qno += 1  # Increment question number for the next pair

                qno += 1  # Increment question number for the next question


    remove_table_borders(combined_table)
    autofit_column_width(combined_table)

    # Footer
    footer_paragraph = document.add_paragraph("\n\n**************")
    footer_paragraph.alignment = 1  # Center alignment
    footer_run = footer_paragraph.runs[0]
    footer_run.font.size = Pt(12)
    footer_run.font.bold = True

    # analysis_path
    all_questions = saqs + laqs
    add_question_analysis_table(document, all_questions)
    # Save document
    document.save(file_path)
